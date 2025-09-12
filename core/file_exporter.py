# core/file_exporter.py
from pathlib import Path
import json
import csv

from config import SAVE_DIR, ALLOWED_FILE_TYPES_OUT

# libs externes
from docx import Document       # pip install python-docx
from fpdf import FPDF           # pip install fpdf2
import openpyxl                 # pip install openpyxl


def export_file(session: str, name: str, content, filetype: str = "txt"):
    """
    Exporte un contenu dans un fichier spécifique.
    - session : nom de la session de chat
    - name : nom de fichier (sans extension)
    - content : texte brut, liste ou dict selon le format
    - filetype : extension cible (txt, md, py, json, csv, docx, pdf, xlsx)
    """
    filetype = filetype.lower()
    if filetype not in ALLOWED_FILE_TYPES_OUT:
        raise ValueError(f"Extension {filetype} non supportée. Autorisées: {ALLOWED_FILE_TYPES_OUT}")

    # chemin cible
    session_dir = Path(SAVE_DIR) / session / "files_out"
    session_dir.mkdir(parents=True, exist_ok=True)
    path = session_dir / f"{name}.{filetype}"

    # === formats simples texte/code ===
    if filetype in ["txt", "py", "md"]:
        path.write_text(str(content), encoding="utf-8")

    # === JSON ===
    elif filetype == "json":
        with open(path, "w", encoding="utf-8") as f:
            json.dump(content, f, ensure_ascii=False, indent=2)

    # === CSV ===
    elif filetype == "csv":
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            if isinstance(content, list):
                for row in content:
                    writer.writerow(row if isinstance(row, (list, tuple)) else [row])
            else:
                f.write(str(content))

    # === DOCX ===
    elif filetype == "docx":
        doc = Document()
        if isinstance(content, list):
            for line in content:
                doc.add_paragraph(str(line))
        else:
            doc.add_paragraph(str(content))
        doc.save(path)

    # === PDF ===
    elif filetype == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        if isinstance(content, list):
            for line in content:
                pdf.multi_cell(0, 10, str(line))
        else:
            pdf.multi_cell(0, 10, str(content))
        pdf.output(str(path))

    # === XLSX ===
    elif filetype == "xlsx":
        wb = openpyxl.Workbook()
        ws = wb.active
        if isinstance(content, list):
            for row in content:
                ws.append(row if isinstance(row, (list, tuple)) else [row])
        else:
            ws.append([str(content)])
        wb.save(path)

    return path
